from tkinter import *
from tkinter import ttk
from tkinter import messagebox
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

import xlrd2
from docx import Document

root = Tk()

workbook = xlrd2.open_workbook('vullist.xlsx')
sheet = workbook.sheet_by_index(0)
row = sheet.nrows
names = sheet.col_values(4)
uyazv_name = sheet.col_values(1)
dates = sheet.col_values(9)
danger_lavels = sheet.col_values(12)
danger_super, danger_hight, danger_middle, danger_low = 0, 0, 0, 0
year_distribution = dict()

list_of_unique_dates_from = []

def get_unique_numbers(default_mas):
    for i in range(8, row):
        default_mas.append(dates[i])

    unique_numbers = sorted(set(default_mas))
    return unique_numbers

root.title('Анализ уязвимостей Windows')
root.geometry("1200x800")

b = Button(root, text="Анализ")
diagramBtn = Button(root, text="Показать диаграмму")
b.place(y=500)
title = Label(root, bg="white", fg='black', font=24)
title['text'] = 'Анализ уязвимостей Windows'
lowLvl = Label(root, bg="white", fg="black", font=24)
mediumLvl = Label(root, bg="white", fg="black", font=24)
highLvl = Label(root, bg="white", fg="black", font=24)
criticalLvl = Label(root, bg="white", fg="black", font=24)
position = {"padx":6, "pady":6, "anchor": NW}
saveDocumentBtn = Button(root, text="Сохранить")

date = "По дате"
default = "За всё время"
options = StringVar(value=default)
dates_list_from = list(get_unique_numbers(list_of_unique_dates_from))

def switch():
    if (options.get() == default):
        combobox_from['state'] = 'disabled'
        combobox_from['values'] = []
    else:
        combobox_from['state'] = 'enabled'
        combobox_from['values'] = dates_list_from

date_btn = ttk.Radiobutton(text=date, value=date, variable=options, command=switch)
default_btn = ttk.Radiobutton(text=default, value=default, variable=options, command=switch)
combobox_from = ttk.Combobox(values=dates_list_from, state="disabled")

title.pack(fill="x", anchor="n", ipady=10)
combobox_from.pack(anchor=NW, padx=6, pady=6)
date_btn.pack(**position)
default_btn.pack(**position)
lowLvl.pack(**position)
mediumLvl.pack(**position)
highLvl.pack(**position)
criticalLvl.pack(**position)
b.pack(pady=5, anchor="s", expand=True)
diagramBtn.pack(pady=5, anchor="s", expand=True)
saveDocumentBtn.pack(pady=5, anchor="s", expand=True)

def saveDocument(e):
    document = Document()

    document.add_heading('Отчет об уязвимостях', 0)
    document.add_paragraph('Низкий уровень опасности - ' + str(danger_low) + '   '  + 'Средний уровень опасности - ' + str(
                danger_middle) + '   '  + 'Высокий уровень опасности - ' + str(danger_hight) + '   '  + 'Критический уровень опасности -' + '   ' + str(
                danger_super))

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Название уязвимости'
    hdr_cells[1].text = 'Год'
    if (options.get() == date and combobox_from.state != 'disabled'):
        for i in range(3, row):
            if (dates[i] == combobox_from.get() and names[i].find('Windows') >= 0):
                row_cells = table.add_row().cells
                row_cells[0].text = uyazv_name[i]
                row_cells[1].text = dates[i]
    if (options.get() == default):
        for i in range(3, row):
            if names[i].find('Windows') >= 0:
                row_cells = table.add_row().cells
                row_cells[0].text = uyazv_name[i]
                row_cells[1].text = dates[i]

    document.add_page_break()
    document.save('demo.docx')

def diagramma(event):
    try:
        if danger_low == 0 and danger_middle == 0 and danger_hight == 0 and danger_super ==0:
            messagebox.showerror('Ошибка',
                                 'Для вывода диаграммы необходимо хотя бы одно значение отличное от 0!')
        else:
            labels = 'Низкий', 'Средний', 'Высокий', 'Критический'
            sizes = [danger_low, danger_middle, danger_hight, danger_super]

            colors = ("grey", "yellow", "orange",
            "brown")
            fig1, ax1 = plt.subplots()
            explode = (0, 0.1, 0, 0)

            ax1.pie(sizes, wedgeprops=dict(width=0.5), colors=colors, explode=explode, labels=labels,autopct='%1.1f%%', shadow=True, startangle=90)
            patches, texts, auto = ax1.pie(sizes, wedgeprops=dict(width=0.5), colors=colors, shadow=True, startangle=90, explode=explode, autopct='%1.1f%%' )

            plt.legend(patches, labels, loc="best")
            okno=Tk()
            okno.title("Диаграмма уязвимостей")
            okno.configure(background='#a8e4a0')
            canvas = FigureCanvasTkAgg(fig1, master=okno)
            canvas.get_tk_widget().pack()
            canvas.draw()
    except NameError:
        messagebox.showerror('Ошибка',
                             'Для вывода диаграммы необходимо провести анализ!')

def choosen(e):
   global danger_super, danger_hight, danger_middle, danger_low
   danger_super = 0
   danger_hight = 0
   danger_middle = 0
   danger_low = 0

   if (options.get() == default):
       for i in range(3, row):
           if names[i].find('Windows') >= 0:
               if danger_lavels[i][0] == 'К':
                   danger_super += 1
               elif danger_lavels[i][0] == 'В':
                   danger_hight += 1
               elif danger_lavels[i][0] == 'С':
                   danger_middle += 1
               else:
                   danger_low += 1
   if (options.get() == date and combobox_from.state != 'disabled'):
       for i in range(3, row):
           if (dates[i] == combobox_from.get() and names[i].find('Windows') >= 0):
               if danger_lavels[i][0] == 'К':
                   danger_super += 1
               elif danger_lavels[i][0] == 'В':
                   danger_hight += 1
               elif danger_lavels[i][0] == 'С':
                   danger_middle += 1
               else:
                   danger_low += 1
   lowLvl['text'] = "Низкого уровня уязвимостей:", danger_low
   mediumLvl['text'] = "Среднего уровня уязвимостей:", danger_middle
   highLvl['text'] = "Высокого уровня уязвимостей:", danger_hight
   criticalLvl['text'] = "Критических уязвимостей:", danger_super
b.bind("<Button-1>", choosen)
date_btn.bind("<<RadioButtonSelected>>", switch)
saveDocumentBtn.bind("<Button-1>", saveDocument)
diagramBtn.bind("<Button-1>", diagramma)

root.update_idletasks()
root.mainloop()